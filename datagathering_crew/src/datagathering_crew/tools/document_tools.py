from crewai.tools import BaseTool
from typing import Type
from pydantic import BaseModel, Field
import os
import json
import docx
from PyPDF2 import PdfReader

class DOCXSearchToolInput(BaseModel):
    """Input schema for DOCXSearchTool."""
    file_path: str = Field(..., description="Path to the DOCX file to search within")
    query: str = Field(..., description="The search query to look for in the file. Use empty string to get all content")

class DOCXSearchTool(BaseTool):
    name: str = "DOCX Search Tool"
    description: str = (
        "A tool for searching and extracting information from DOCX resume files. "
        "Provide the file path and your search query to extract relevant information. "
        "Use an empty string as query to get all content from the document."
    )
    args_schema: Type[BaseModel] = DOCXSearchToolInput

    def _run(self, file_path: str, query: str) -> str:
        result = {
            "status": "success",
            "content": "",
            "metadata": {
                "file_path": file_path,
                "query": query
            }
        }
        
        try:
            # Check for invalid file path inputs
            if not file_path or not isinstance(file_path, str) or file_path.startswith("File not found") or file_path == "No resume file provided":
                result["status"] = "error"
                result["message"] = f"No valid resume file provided. Input was: {file_path}"
                return json.dumps(result, indent=2)
                
            if not os.path.exists(file_path):
                result["status"] = "error"
                result["message"] = f"File not found at path '{file_path}'. Please check if the file exists and the path is correct."
                return json.dumps(result, indent=2)
            
            # Check file extension
            if not file_path.lower().endswith('.docx'):
                result["status"] = "error"
                result["message"] = f"The file at '{file_path}' is not a DOCX file. Please provide a file with .docx extension."
                return json.dumps(result, indent=2)
            
            doc = docx.Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            if not content.strip():
                result["status"] = "warning"
                result["message"] = f"The document at '{file_path}' appears to be empty or contains no readable text."
                return json.dumps(result, indent=2)
            
            # If query is empty, return all content
            if not query or query.strip() == "":
                result["content"] = content
                return json.dumps(result, indent=2)
            
            # Simple search implementation
            query = query.lower()
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and query in text.lower():
                    paragraphs.append(text)
            
            if not paragraphs:
                result["status"] = "no_match"
                result["message"] = f"No content exactly matching '{query}' found in the document."
                result["content"] = content  # Return full content for general processing
                return json.dumps(result, indent=2)
            
            result["content"] = "\n\n".join(paragraphs)
            return json.dumps(result, indent=2)
        
        except Exception as e:
            result["status"] = "error"
            result["message"] = f"Error processing DOCX file '{file_path}': {str(e)}"
            return json.dumps(result, indent=2)

class PDFSearchToolInput(BaseModel):
    """Input schema for PDFSearchTool."""
    file_path: str = Field(..., description="Path to the PDF file to search within")
    query: str = Field(..., description="The search query to look for in the file. Use empty string to get all content")

class PDFSearchTool(BaseTool):
    name: str = "PDF Search Tool"
    description: str = (
        "A tool for searching and extracting information from PDF resume files. "
        "Provide the file path and your search query to extract relevant information. "
        "Use an empty string as query to get all content from the document."
    )
    args_schema: Type[BaseModel] = PDFSearchToolInput

    def _run(self, file_path: str, query: str) -> str:
        result = {
            "status": "success",
            "content": "",
            "metadata": {
                "file_path": file_path,
                "query": query
            }
        }
        
        try:
            # Check for invalid file path inputs
            if not file_path or not isinstance(file_path, str) or file_path.startswith("File not found") or file_path == "No resume file provided":
                result["status"] = "error"
                result["message"] = f"No valid resume file provided. Input was: {file_path}"
                return json.dumps(result, indent=2)
                
            if not os.path.exists(file_path):
                result["status"] = "error"
                result["message"] = f"File not found at path '{file_path}'. Please check if the file exists and the path is correct."
                return json.dumps(result, indent=2)
            
            # Check file extension
            if not file_path.lower().endswith('.pdf'):
                result["status"] = "error"
                result["message"] = f"The file at '{file_path}' is not a PDF file. Please provide a file with .pdf extension."
                return json.dumps(result, indent=2)
            
            try:
                reader = PdfReader(file_path)
            except Exception as pdf_error:
                result["status"] = "error"
                result["message"] = f"Error opening PDF file: {str(pdf_error)}"
                return json.dumps(result, indent=2)
            
            if len(reader.pages) == 0:
                result["status"] = "warning"
                result["message"] = f"The PDF at '{file_path}' appears to be empty or contains no pages."
                return json.dumps(result, indent=2)
            
            content = ""
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
                except Exception as extract_error:
                    # Continue even if one page fails
                    pass
            
            if not content.strip():
                result["status"] = "warning"
                result["message"] = f"The PDF at '{file_path}' contains no extractable text. It might be a scanned document or image-based PDF."
                return json.dumps(result, indent=2)
            
            # If query is empty, return all content
            if not query or query.strip() == "":
                result["content"] = content
                return json.dumps(result, indent=2)
            
            # Simple search implementation
            query = query.lower()
            content_lower = content.lower()
            
            if query not in content_lower:
                result["status"] = "no_match"
                result["message"] = f"No content exactly matching '{query}' found in the document."
                result["content"] = content  # Return full content for general processing
                return json.dumps(result, indent=2)
            
            # Extract paragraphs that contain the query
            paragraphs = content.split('\n\n')
            matching_paragraphs = [p for p in paragraphs if query in p.lower()]
            
            if not matching_paragraphs:
                # Fall back to returning a portion around the query
                idx = content_lower.find(query)
                start = max(0, content_lower.rfind('\n', 0, max(0, idx - 200)))
                end = content_lower.find('\n', min(len(content_lower), idx + len(query) + 200))
                if end == -1:
                    end = len(content_lower)
                result["content"] = content[start:end]
            else:
                result["content"] = "\n\n".join(matching_paragraphs)
                
            return json.dumps(result, indent=2)
        
        except Exception as e:
            result["status"] = "error"
            result["message"] = f"Error processing PDF file '{file_path}': {str(e)}"
            return json.dumps(result, indent=2) 